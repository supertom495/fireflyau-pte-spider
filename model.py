import requests
import json
import time 
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import model
import re

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    r.font.underline = True

    return hyperlink


def writeHeader(result, document, n):

    questionNumStr = result["questionNumStr"]
    title = result["title"]

    document.add_heading("第" + str(n) + "篇 #" + questionNumStr, level=4)
    document.add_heading(title, level=1)
    document.add_paragraph("").style.font.name = u'Cambria (Body)'



def writeSST(result, document):
    answerInfo = result["answerInfo"]               # 缩减答案
    answerTranscript = result["answerTranscript"]   # 原文
    idNum = result["id"]                            # uuid4
    question = result["question"]                   # audio link
    videoUrl = result["videoUrl"]                   # 好像没用
    

    if answerTranscript != "":
        document.add_paragraph(answerTranscript + "\n")
    else :
        document.add_paragraph("无原文\n")

    document.add_paragraph("参考答案: ")
    document.add_paragraph(answerInfo + "\n")


    if question != "":
        p = document.add_paragraph("音频")
        add_hyperlink(p, 'audio link', question)
    else :
        document.add_paragraph("无音频")

    document.add_page_break()



def writeRWFIB(result, document):
    answerInfo = result["answerInfo"]               # 答案注释
    idNum = result["id"]                            # uuid4 没用
    questionInfo = result["questionInfo"]           # 原文
    videoUrl = result["videoUrl"]                   # 好像没用


    splitedLst = re.compile(r"\[(.*?)\]").split(questionInfo)
    # print(len(splitedLst))
    
    p = document.add_paragraph(None)

    if questionInfo != "":
        count = 0
        for item in splitedLst:
            if count%2 != 0:        # set the every second item to bold and with color
                a = p.add_run(item)   
                a.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
                a.bold = True
            else:
                p.add_run(item)
            count += 1
    else :
        document.add_paragraph("无原文\n")

    document.add_paragraph("注释: ")
    document.add_paragraph(answerInfo)

    if videoUrl != "":
        p = document.add_paragraph("Some Video")
        add_hyperlink(p, 'link', videoUrl)


    document.add_page_break()


def writeResult(result, document, n, name):

    writeHeader(result, document, n)

    if name == "LSST":
        writeSST(result, document)
    
    if name == "RWFIB":
        writeRWFIB(result, document)

    if name == "RFIB":
        writeRWFIB(result, document)
    
    
    
    
    print("第" + str(n) + "篇"+ name + " #" + result["questionNumStr"])
