import requests
import json
import time 
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.dml import MSO_THEME_COLOR_INDEX



document = Document()

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    r.font.underline = True

    return hyperlink


def postKey(n):
    url = "https://api.fireflyau.com/yhcapp_n/pte/other/findOneQuestionPrediction"

    # name = base64.b64decode(b'WGl5YW4gVG9uZw==').decode("utf-8")
    # email = base64.b64decode(b'eGl5YW4udG9uZ0BvdXRsb29rLmNvbQ==').decode("utf-8")
    
    jsona = {"bark":1,"num":n,"button":3,"questionType":"LSST","isJJ":1,"submitNum":1,"pId":"8dd66b6b6348c45180622b5b3522688f"}

    info = {
        "data": str(jsona),
        "token": "620b6957b18a6bf7e99accf51f1c108f_d1cb0d8c45bca2f4006c0a18dae2516a"
    }


    # data = json.dumps(info)

    response = requests.post(url, data=info)

    res = json.loads(response.text)
    print(res)
    print("状态： " + res["status"] + "第" + str(n+1) + "篇")
    # with open(r"API_data\log.txt", 'a+') as f:
    #     f.write(str(datetime.datetime.now()))
    #     f.write("\n")
    #     f.write(str(message))
    #     f.write("\n\n")


    data = res["data"]
    question = data["question"]

    title = question["title"]
    sampleAnswer = question["answerInfo"]
    original = question["answerTranscript"]
    audio = question["question"]
    questionNumStr = question["questionNumStr"]

    addInfo(title, original, sampleAnswer, audio, n, questionNumStr)



def addInfo(title, original, sample, url, n, number):
    global document

    document.add_heading("第" + str(n+1) + "篇 #" + number, level=4)
    document.add_heading(title, level=1)

    document.add_paragraph("")


    if original != "":
        document.add_paragraph(original)
    else :
        document.add_paragraph("无原文")

    document.add_paragraph("")
    document.add_paragraph("参考答案: ")
    document.add_paragraph(sample)
    document.add_paragraph("") 


    if url != "":
        p = document.add_paragraph("音频")
        add_hyperlink(p, 'audio link', url)
    else :
        document.add_paragraph("无音频")


    document.add_page_break()



if __name__ == "__main__":
    # global document 
    document = Document()
    document.add_heading('SST 5.2', 0)

    for i in range(0, 66): 
        # print(i)
        postKey(i)
        time.sleep(0.5)
    # postKey(0)
    document.save('SST 5.2.docx')



