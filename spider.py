import requests
import json
import time 
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import model



def getSetting():
    with open(r"setting/setting.json", 'r') as f:
        setting = json.load(f)
    return setting["url"], setting["token"]


def getData():
    with open(r"setting/data.json", 'r') as f:
        setting = json.load(f)
    return setting


def getResponse(name, num):
    url, token = getSetting()
    data = getData()
    data["questionType"] = name
    data["num"] = num
    form = {
        "token": token,
        "data": json.dumps(data)
    }

    response = requests.post(url, data=form)
    
    if response.status_code != 200:
        time.sleep(0.5)
        return getResponse(name, num)

    return response



    


def logError(name, result, num):
    with open(r"file/errorLog.txt", "a+") as f:
        f.write(name + " 第" + str(num) + "篇没拿到\n")
        f.write("原因: " + str(result["ErrorInfo"]))
        f.write("\n\n")

def writeToDocument(name, document):
    result = getResponse(name, 1) # try to get the first result
    # if result.status_code != "200":
    lastNum = int(json.loads(result.text)["data"]["numCount"]) + 1

    for i in range(0, lastNum):
        result = json.loads(getResponse(name, i).text)

        if result["status"] == "1":
            model.writeResult(result["data"]["question"], document, i+1, name)
        else: 
            logError(name, result, i)

        time.sleep(1)



def run(name, version):
    document = Document()
    fileName = name + version
    # print(fileName)

    document.styles['Normal'].font.name = u'SimHei'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'SimHei')

    document.add_heading(fileName, 0)
    writeToDocument(name, document)

    document.save("file/" + fileName + ".docx")


if __name__ == "__smain__":
    # global document 
    document = Document()
    documentType = "LSST"
    documentVersion = "5.2"

    document.add_heading(documentType + documentVersion, 0)
    
    for i in range(0, 66): 
        # print(i)
        postKey(i)
        time.sleep(0.5)
    # postKey(0)
    document.save(""+ documentType + documentVersion + '.docx')



